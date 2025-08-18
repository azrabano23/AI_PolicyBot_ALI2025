#!/usr/bin/env python3
"""
CSRR Faculty Media Tracker
Automated system to track and report on faculty affiliate publications
"""

import os
import json
import schedule
import time
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime, timedelta
import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
from flask import Flask, render_template, request, jsonify
import threading
from pathlib import Path
import re
import logging
from typing import Dict, List, Optional
import yaml

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class MediaTracker:
    def __init__(self, config_path="config.yaml"):
        """Initialize the media tracker with configuration"""
        self.config = self.load_config(config_path)
        self.data_file = "media_data.json"
        self.faculty_list = self.load_faculty_list()
        self.search_queries = self.generate_search_queries()
        
    def load_config(self, config_path):
        """Load configuration from YAML file"""
        try:
            with open(config_path, 'r') as file:
                return yaml.safe_load(file)
        except FileNotFoundError:
            # Default configuration if file doesn't exist
            return {
                'email': {
                    'smtp_server': 'smtp.gmail.com',
                    'smtp_port': 587,
                    'sender_email': 'your_email@gmail.com',
                    'sender_password': 'your_password',
                    'recipient_email': 'boss@csrr.rutgers.edu'
                },
                'search_sources': [
                    'google_news',
                    'bing_news',
                    'twitter',
                    'academic_databases'
                ],
                'keywords': [
                    'op-ed', 'opinion', 'interview', 'television',
                    'radio', 'podcast', 'commentary', 'analysis'
                ]
            }
    
    def load_faculty_list(self):
        """Load faculty affiliate list from the original document"""
        # This list is extracted from the document you provided
        faculty_list = [
            "Zain Abdullah", "Matthew Abraham", "Atiya Aftab", "Ghada Ageel",
            "Nadia Ahmad", "Aziza Ahmed", "Susan M. Akram", "M. Shahid Alam",
            "Khalil Al-Anani", "Raquel E Aldana", "Omar Al-Dewachi", "Tazeen M. Ali",
            "Zahra Ali", "Ousseina Alidou", "Sabrina Alimahomed-Wilson", "Nermin Allam",
            "Mohamed Alsiadi", "Mason Ameri", "Leyla Amzi-Erdogdular", "Mohamed 'Arafa",
            "Abed Awad", "Muhannad Ayyash", "Gaiutra Devi Bahadur", "Asli Ü. Bâli",
            "William C. Banks", "Esther Canty-Barnes", "Beth Baron", "Hatem Bazian",
            "Rabea Benhalim", "Emily Berman", "Khaled A. Beydoun", "George Bisharat",
            "Bidisha Biswas", "Elise Boddie", "Mark Bray", "Umayyah Cable",
            "Robert S. Chang", "Ali R. Chaudhary", "Cyra A. Choudhury", "LaToya Baldwin Clark",
            "Juan Cole", "Jorge Contesse", "Omar S. Dahi", "Omar Dajani",
            "Karam Dana", "Timothy P. Daniels", "Meera E. Deo", "Karishma Desai",
            "Veena Dubal", "Jon Dubin", "Stephen Dycus", "Timothy Eatman",
            "Taleed El-Sabawi", "Sarah Eltantawi", "Noura Erakat", "John L. Esposito",
            "Marta Esquilin", "Mohammad Fadel", "Dalia Fahmy", "Huda J. Fakhreddine",
            "John Farmer, Jr.", "Jonathan Feingold", "Katherine M. Franke", "Brittany Friedman",
            "Emmaia Gelman", "Ameena Ghaffar-Kucher", "Behrooz Ghamari-Tabrizi",
            "D. Asher Ghertner", "Rachel Godsil", "Wendy Greene", "Catherine M. Grosso",
            "Anju Gupta", "Zeynep Devrim Gürsel", "Farid Hafez", "Jonathan Hafetz",
            "Haider Ala Hamoudi", "Rebecca Hankins", "Adil Haque"
        ]
        return faculty_list
    
    def generate_search_queries(self):
        """Generate search queries for each faculty member"""
        queries = []
        for faculty in self.faculty_list:
            for keyword in self.config['keywords']:
                queries.append(f'"{faculty}" {keyword}')
        return queries
    
    def search_google_news(self, query: str, days_back: int = 30) -> List[Dict]:
        """Search Google News for faculty mentions"""
        results = []
        try:
            # This is a simplified example - you'd need to use Google News API
            # or a web scraping approach with proper rate limiting
            url = f"https://news.google.com/search?q={query}&hl=en-US&gl=US&ceid=US%3Aen"
            
            # Note: In production, you'd need to use official APIs or handle rate limiting
            # This is a placeholder for the actual implementation
            logger.info(f"Searching Google News for: {query}")
            
            # Placeholder result structure
            result = {
                'title': f"Sample article for {query}",
                'url': 'https://example.com/article',
                'publication': 'Sample Publication',
                'date': datetime.now().strftime('%Y-%m-%d'),
                'faculty_name': query.split('"')[1],
                'type': 'news_article'
            }
            results.append(result)
            
        except Exception as e:
            logger.error(f"Error searching Google News: {e}")
        
        return results
    
    def search_academic_databases(self, faculty_name: str) -> List[Dict]:
        """Search academic databases for faculty publications"""
        results = []
        # This would integrate with academic databases like JSTOR, Project MUSE, etc.
        # Placeholder implementation
        logger.info(f"Searching academic databases for: {faculty_name}")
        return results
    
    def search_social_media(self, faculty_name: str) -> List[Dict]:
        """Search social media for faculty mentions"""
        results = []
        # This would integrate with Twitter API, etc.
        # Placeholder implementation
        logger.info(f"Searching social media for: {faculty_name}")
        return results
    
    def run_daily_search(self):
        """Run daily search for all faculty members"""
        logger.info("Starting daily search...")
        all_results = []
        
        for faculty in self.faculty_list:
            # Search different sources
            news_results = self.search_google_news(f'"{faculty}"')
            academic_results = self.search_academic_databases(faculty)
            social_results = self.search_social_media(faculty)
            
            all_results.extend(news_results)
            all_results.extend(academic_results)
            all_results.extend(social_results)
        
        # Save results
        self.save_results(all_results)
        logger.info(f"Daily search completed. Found {len(all_results)} results.")
        
        return all_results
    
    def save_results(self, results: List[Dict]):
        """Save search results to JSON file"""
        try:
            # Load existing data
            if os.path.exists(self.data_file):
                with open(self.data_file, 'r') as f:
                    existing_data = json.load(f)
            else:
                existing_data = []
            
            # Add new results
            existing_data.extend(results)
            
            # Remove duplicates based on URL
            unique_results = []
            seen_urls = set()
            for result in existing_data:
                if result['url'] not in seen_urls:
                    unique_results.append(result)
                    seen_urls.add(result['url'])
            
            # Save updated data
            with open(self.data_file, 'w') as f:
                json.dump(unique_results, f, indent=2)
                
        except Exception as e:
            logger.error(f"Error saving results: {e}")
    
    def generate_monthly_report(self):
        """Generate monthly report"""
        try:
            # Load data
            with open(self.data_file, 'r') as f:
                data = json.load(f)
            
            # Filter for last month
            last_month = datetime.now() - timedelta(days=30)
            recent_data = [
                item for item in data 
                if datetime.strptime(item['date'], '%Y-%m-%d') >= last_month
            ]
            
            # Generate Excel report
            excel_file = self.create_excel_report(recent_data)
            
            # Generate Word report
            word_file = self.create_word_report(recent_data)
            
            # Send email with reports
            self.send_monthly_email(excel_file, word_file, recent_data)
            
            logger.info("Monthly report generated and sent.")
            
        except Exception as e:
            logger.error(f"Error generating monthly report: {e}")
    
    def create_excel_report(self, data: List[Dict]) -> str:
        """Create Excel report from data"""
        df = pd.DataFrame(data)
        
        # Group by faculty member
        faculty_groups = df.groupby('faculty_name')
        
        filename = f"csrr_media_report_{datetime.now().strftime('%Y_%m')}.xlsx"
        
        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            # Summary sheet
            summary_data = []
            for faculty, group in faculty_groups:
                summary_data.append({
                    'Faculty Name': faculty,
                    'Total Articles': len(group),
                    'Op-Eds': len(group[group['type'] == 'op-ed']),
                    'Interviews': len(group[group['type'] == 'interview']),
                    'Other': len(group[group['type'] == 'other'])
                })
            
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
            
            # Individual faculty sheets
            for faculty, group in faculty_groups:
                sheet_name = faculty.replace(' ', '_')[:31]  # Excel sheet name limit
                group.to_excel(writer, sheet_name=sheet_name, index=False)
        
        return filename
    
    def create_word_report(self, data: List[Dict]) -> str:
        """Create Word document report"""
        doc = Document()
        
        # Title
        title = doc.add_heading('CSRR Faculty Affiliates Media Report', 0)
        title.alignment = 1  # Center alignment
        
        # Date
        doc.add_paragraph(f"Report Period: {datetime.now().strftime('%B %Y')}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(f"Total articles found: {len(data)}")
        doc.add_paragraph(f"Number of faculty with media coverage: {len(set(item['faculty_name'] for item in data))}")
        
        # By faculty member
        doc.add_heading('Faculty Media Coverage', level=1)
        
        faculty_dict = {}
        for item in data:
            faculty_name = item['faculty_name']
            if faculty_name not in faculty_dict:
                faculty_dict[faculty_name] = []
            faculty_dict[faculty_name].append(item)
        
        for faculty, articles in faculty_dict.items():
            doc.add_heading(faculty, level=2)
            
            for article in articles:
                p = doc.add_paragraph()
                p.add_run(article['title']).bold = True
                p.add_run(f", {article['publication']}, {article['date']}")
                p.add_run(f"\n{article['url']}")
        
        filename = f"csrr_media_report_{datetime.now().strftime('%Y_%m')}.docx"
        doc.save(filename)
        
        return filename
    
    def send_monthly_email(self, excel_file: str, word_file: str, data: List[Dict]):
        """Send monthly email report"""
        try:
            msg = MIMEMultipart()
            msg['From'] = self.config['email']['sender_email']
            msg['To'] = self.config['email']['recipient_email']
            msg['Subject'] = f"CSRR Faculty Media Report - {datetime.now().strftime('%B %Y')}"
            
            # Email body
            body = f"""
            Dear CSRR Team,
            
            Please find attached the monthly faculty media report for {datetime.now().strftime('%B %Y')}.
            
            Summary:
            - Total articles found: {len(data)}
            - Faculty members with coverage: {len(set(item['faculty_name'] for item in data))}
            
            The report includes:
            - Excel spreadsheet with detailed data
            - Word document with formatted report
            
            This report was automatically generated by the CSRR Media Tracker system.
            
            Best regards,
            CSRR Media Tracker Bot
            """
            
            msg.attach(MIMEText(body, 'plain'))
            
            # Attach Excel file
            with open(excel_file, 'rb') as attachment:
                part = MIMEMultipart()
                part.set_payload(attachment.read())
                part.add_header('Content-Disposition', f'attachment; filename= {excel_file}')
                msg.attach(part)
            
            # Attach Word file
            with open(word_file, 'rb') as attachment:
                part = MIMEMultipart()
                part.set_payload(attachment.read())
                part.add_header('Content-Disposition', f'attachment; filename= {word_file}')
                msg.attach(part)
            
            # Send email
            server = smtplib.SMTP(self.config['email']['smtp_server'], self.config['email']['smtp_port'])
            server.starttls()
            server.login(self.config['email']['sender_email'], self.config['email']['sender_password'])
            server.send_message(msg)
            server.quit()
            
            logger.info("Monthly email sent successfully.")
            
        except Exception as e:
            logger.error(f"Error sending email: {e}")

# Flask Web Application
app = Flask(__name__)
tracker = MediaTracker()

@app.route('/')
def dashboard():
    """Main dashboard"""
    return render_template('dashboard.html')

@app.route('/api/data')
def get_data():
    """API endpoint to get media data"""
    try:
        if os.path.exists(tracker.data_file):
            with open(tracker.data_file, 'r') as f:
                data = json.load(f)
        else:
            data = []
        
        return jsonify(data)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/search', methods=['POST'])
def manual_search():
    """API endpoint for manual search"""
    try:
        results = tracker.run_daily_search()
        return jsonify({'status': 'success', 'results_count': len(results)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/report', methods=['POST'])
def generate_report():
    """API endpoint to generate report"""
    try:
        tracker.generate_monthly_report()
        return jsonify({'status': 'success', 'message': 'Report generated and sent'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def run_scheduler():
    """Run the scheduler in a separate thread"""
    # Schedule daily searches
    schedule.every().day.at("09:00").do(tracker.run_daily_search)
    
    # Schedule monthly report on 1st of each month
    schedule.every().month.do(tracker.generate_monthly_report)
    
    while True:
        schedule.run_pending()
        time.sleep(3600)  # Check every hour

if __name__ == '__main__':
    # Start scheduler in background
    scheduler_thread = threading.Thread(target=run_scheduler)
    scheduler_thread.daemon = True
    scheduler_thread.start()
    
    # Run Flask app
    app.run(debug=True, host='0.0.0.0', port=5000)
